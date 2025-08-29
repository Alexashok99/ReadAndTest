import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter.scrolledtext import ScrolledText
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from docx import Document
from pathlib import Path

# --------- Helpers ---------
def extract_text_from_docx(path: Path) -> str:
    """
    .docx से सभी paragraphs और tables का टेक्स्ट निकालता है।
    """
    doc = Document(str(path))

    parts = []

    # Paragraphs
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            parts.append(txt)

    # Tables (row-wise join)
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            row_line = " | ".join([c for c in cells if c])
            if row_line:
                parts.append(row_line)

    # Sections जैसे header/footer आमतौर पर python-docx से सीधे नहीं मिलते
    return "\n".join(parts).strip()


def make_viewer_window(master: ttk.Window, title: str, content: str):
    """
    .docx का कंटेंट एक नई Toplevel विंडो में read-only ScrolledText में दिखाता है।
    """
    win = ttk.Toplevel(master)
    win.title(title)
    win.geometry("900x600")
    win.place_window_center()

    # Top bar info
    info = ttk.Label(win, text=title, bootstyle=INFO, anchor="w")
    info.pack(fill="x", padx=8, pady=(8, 4))

    # ScrolledText (read-only)
    txt = ScrolledText(win, wrap="word", font=("Consolas", 11))
    txt.pack(fill="both", expand=True, padx=8, pady=8)

    # Insert content
    txt.configure(state="normal")
    if content:
        txt.insert("1.0", content)
    else:
        txt.insert("1.0", "[Empty document]")
    txt.configure(state="disabled")

    # Quick find box (optional nicety)
    find_frame = ttk.Frame(win)
    find_frame.pack(fill="x", padx=8, pady=(0, 8))
    ttk.Label(find_frame, text="Find:").pack(side="left")
    find_var = tk.StringVar()
    find_entry = ttk.Entry(find_frame, textvariable=find_var, width=30)
    find_entry.pack(side="left", padx=(6, 6))

    def do_find(event=None):
        query = find_var.get().strip()
        txt.tag_remove("match", "1.0", "end")
        if not query:
            return
        start = "1.0"
        while True:
            pos = txt.search(query, start, stopindex="end", nocase=True)
            if not pos:
                break
            end = f"{pos}+{len(query)}c"
            txt.tag_add("match", pos, end)
            start = end
        txt.tag_config("match", background="yellow")
        # Scroll to first match
        first = txt.tag_ranges("match")
        if first:
            txt.see(first[0])

    ttk.Button(find_frame, text="Search", command=do_find, bootstyle=SECONDARY).pack(side="left")
    find_entry.bind("<Return>", do_find)

    # Allow mouse wheel scrolling on Windows/Linux/Mac
    def _on_mousewheel(event):
        # normalize delta across platforms
        delta = 0
        if event.num == 4:   # Linux scroll up
            delta = -1
        elif event.num == 5: # Linux scroll down
            delta = 1
        elif event.delta:
            delta = -1 if event.delta > 0 else 1
        txt.yview_scroll(delta, "units")
        return "break"

    txt.bind("<MouseWheel>", _on_mousewheel)   # Windows / macOS
    txt.bind("<Button-4>", _on_mousewheel)     # Linux
    txt.bind("<Button-5>", _on_mousewheel)     # Linux


# --------- Main App ---------
class DocxReaderApp(ttk.Window):
    def __init__(self, themename="flatly"):
        super().__init__(themename=themename)
        self.title("DOCX Reader - ttkbootstrap")
        self.geometry("700x420")
        self.place_window_center()

        self._build_ui()
        self._build_menubar()
        self._bind_shortcuts()

    def _build_ui(self):
        # Header
        header = ttk.Label(self, text="Open a .docx file to preview",
                           bootstyle=PRIMARY, font=("Segoe UI", 14, "bold"))
        header.pack(fill="x", padx=12, pady=(12, 8))

        # Card
        card = ttk.Frame(self, padding=16, bootstyle="secondary")
        card.pack(fill="both", expand=True, padx=12, pady=12)

        # Instruction
        ttk.Label(card, text="Click the button below or use Ctrl+O to choose a .docx file.",
                  anchor="w").pack(fill="x", pady=(0, 12))

        # Open Button
        ttk.Button(card, text="Open .docx", bootstyle=SUCCESS, command=self.open_docx)\
            .pack(pady=(0, 8))

        # Theme switcher (nice extra)
        themes = ttk.Style().theme_names()
        self.theme_var = tk.StringVar(value=self.style.theme.name)
        theme_row = ttk.Frame(card)
        theme_row.pack(pady=(16, 0))
        ttk.Label(theme_row, text="Theme:").pack(side="left")
        theme_cb = ttk.Combobox(theme_row, values=themes, textvariable=self.theme_var, width=20, state="readonly")
        theme_cb.pack(side="left", padx=8)

        def change_theme(event=None):
            new_theme = self.theme_var.get()
            try:
                self.style.theme_use(new_theme)
            except Exception as e:
                messagebox.showerror("Theme Error", str(e))

        theme_cb.bind("<<ComboboxSelected>>", change_theme)

        # Footer
        footer = ttk.Label(self, text="© DOCX Reader Demo • ttkbootstrap", anchor="center",
                           bootstyle="secondary")
        footer.pack(fill="x", padx=12, pady=(0, 10))

    def _build_menubar(self):
        menubar = tk.Menu(self)
        filemenu = tk.Menu(menubar, tearoff=0)
        filemenu.add_command(label="Open… (Ctrl+O)", command=self.open_docx)
        filemenu.add_separator()
        filemenu.add_command(label="Exit (Ctrl+Q)", command=self.destroy)
        menubar.add_cascade(label="File", menu=filemenu)
        self.config(menu=menubar)

    def _bind_shortcuts(self):
        self.bind_all("<Control-o>", lambda e: self.open_docx())
        self.bind_all("<Control-q>", lambda e: self.destroy())

    def open_docx(self):
        filetypes = [("Word Document", "*.docx"), ("All files", "*.*")]
        path_str = filedialog.askopenfilename(title="Select a .docx file", filetypes=filetypes)
        if not path_str:
            return
        path = Path(path_str)

        if path.suffix.lower() != ".docx":
            messagebox.showwarning("Unsupported", "Please select a .docx file.")
            return

        try:
            content = extract_text_from_docx(path)
        except Exception as e:
            messagebox.showerror("Read Error", f"Failed to read file:\n{e}")
            return

        # Small metadata line
        title = f"{path.name}  —  {path.stat().st_size} bytes"
        make_viewer_window(self, title, content)


if __name__ == "__main__":
    app = DocxReaderApp(themename="flatly")  # you can change theme here
    app.mainloop()
